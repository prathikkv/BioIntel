import os
import json
import tempfile
import time
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path
import io
import base64

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from jinja2 import Environment, FileSystemLoader, Template
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    print("WeasyPrint not available - PDF generation disabled")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available - DOCX generation disabled")
from fastapi import HTTPException, status

from models.database import get_db
from models.bioinformatics import Dataset, AnalysisJob, ExpressionData
from models.literature import LiteratureSummary, Report
from models.user import User
from utils.logging import get_logger
from utils.config import get_settings

logger = get_logger(__name__)
settings = get_settings()

class ReportsService:
    """Service for generating comprehensive reports"""
    
    def __init__(self):
        self.db = next(get_db())
        self.templates_dir = Path(__file__).parent.parent / "templates" / "reports"
        self.reports_dir = Path(settings.REPORTS_DIR) if hasattr(settings, 'REPORTS_DIR') else Path(tempfile.gettempdir()) / "biointel_reports"
        
        # Create directories if they don't exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
    
    @staticmethod
    async def initialize():
        """Initialize reports service"""
        logger.info("Reports service initialized")
    
    async def generate_report(self, user_id: int, report_request: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a comprehensive report"""
        start_time = time.time()
        
        try:
            # Collect data based on report type
            report_data = await self._collect_report_data(user_id, report_request)
            
            # Generate report content
            content = await self._generate_report_content(report_data, report_request)
            
            # Create report record
            report = Report(
                user_id=user_id,
                report_type=report_request["report_type"],
                title=report_request["title"],
                description=report_request.get("description"),
                content=content,
                dataset_ids=report_request.get("dataset_ids"),
                analysis_job_ids=report_request.get("analysis_job_ids"),
                literature_summary_ids=report_request.get("literature_summary_ids"),
                template_used=report_request.get("template_name", "default"),
                format_type=report_request.get("format_type", "html")
            )
            
            self.db.add(report)
            self.db.commit()
            self.db.refresh(report)
            
            # Generate file if requested format is not HTML
            if report.format_type != "html":
                file_path = await self._generate_report_file(report)
                report.file_path = file_path
                report.file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
            
            # Update generation time
            generation_time = time.time() - start_time
            report.generation_time = generation_time
            
            self.db.commit()
            
            logger.info(f"Report generated successfully: {report.id}")
            
            return report.to_dict()
            
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Report generation failed: {str(e)}"
            )
    
    async def _collect_report_data(self, user_id: int, report_request: Dict[str, Any]) -> Dict[str, Any]:
        """Collect all necessary data for report generation"""
        data = {
            "user_info": self._get_user_info(user_id),
            "report_metadata": {
                "title": report_request["title"],
                "description": report_request.get("description"),
                "report_type": report_request["report_type"],
                "generated_at": datetime.utcnow().isoformat(),
                "template": report_request.get("template_name", "default")
            },
            "datasets": [],
            "analysis_results": [],
            "literature_summaries": [],
            "statistics": {},
            "plots": []
        }
        
        # Collect dataset information
        if report_request.get("dataset_ids"):
            for dataset_id in report_request["dataset_ids"]:
                dataset_info = await self._get_dataset_info(dataset_id, user_id)
                if dataset_info:
                    data["datasets"].append(dataset_info)
        
        # Collect analysis results
        if report_request.get("analysis_job_ids"):
            for job_id in report_request["analysis_job_ids"]:
                analysis_info = await self._get_analysis_info(job_id, user_id)
                if analysis_info:
                    data["analysis_results"].append(analysis_info)
        
        # Collect literature summaries
        if report_request.get("literature_summary_ids"):
            for summary_id in report_request["literature_summary_ids"]:
                literature_info = await self._get_literature_info(summary_id, user_id)
                if literature_info:
                    data["literature_summaries"].append(literature_info)
        
        # Calculate summary statistics
        data["statistics"] = self._calculate_summary_statistics(data)
        
        return data
    
    def _get_user_info(self, user_id: int) -> Dict[str, Any]:
        """Get user information"""
        user = self.db.query(User).filter(User.id == user_id).first()
        if user:
            return {
                "name": user.full_name or "Unknown User",
                "email": user.email,
                "organization": user.organization or "Unknown Organization"
            }
        return {"name": "Unknown User", "email": "unknown@example.com", "organization": "Unknown Organization"}
    
    async def _get_dataset_info(self, dataset_id: int, user_id: int) -> Optional[Dict[str, Any]]:
        """Get dataset information"""
        dataset = self.db.query(Dataset).filter(
            Dataset.id == dataset_id,
            Dataset.user_id == user_id
        ).first()
        
        if not dataset:
            return None
        
        # Get basic dataset info
        dataset_info = dataset.to_dict()
        
        # Add quality metrics
        dataset_info["quality_metrics"] = {
            "quality_score": dataset.data_quality_score,
            "missing_values": dataset.missing_values_count,
            "completeness": ((dataset.num_genes * dataset.num_samples - dataset.missing_values_count) / (dataset.num_genes * dataset.num_samples)) * 100 if dataset.num_genes and dataset.num_samples else 0
        }
        
        return dataset_info
    
    async def _get_analysis_info(self, job_id: int, user_id: int) -> Optional[Dict[str, Any]]:
        """Get analysis job information"""
        job = self.db.query(AnalysisJob).filter(
            AnalysisJob.id == job_id,
            AnalysisJob.user_id == user_id
        ).first()
        
        if not job:
            return None
        
        analysis_info = job.to_dict()
        
        # Add interpretation based on job type
        if job.job_type == "pca":
            analysis_info["interpretation"] = self._interpret_pca_results(job.results)
        elif job.job_type == "clustering":
            analysis_info["interpretation"] = self._interpret_clustering_results(job.results)
        
        return analysis_info
    
    async def _get_literature_info(self, summary_id: int, user_id: int) -> Optional[Dict[str, Any]]:
        """Get literature summary information"""
        summary = self.db.query(LiteratureSummary).filter(
            LiteratureSummary.id == summary_id,
            LiteratureSummary.user_id == user_id
        ).first()
        
        if not summary:
            return None
        
        return summary.to_dict()
    
    def _interpret_pca_results(self, results: Dict[str, Any]) -> str:
        """Interpret PCA results"""
        if not results or "explained_variance" not in results:
            return "PCA analysis completed but results interpretation unavailable."
        
        explained_variance = results["explained_variance"]
        if len(explained_variance) >= 2:
            pc1_var = explained_variance[0] * 100
            pc2_var = explained_variance[1] * 100
            total_var = (explained_variance[0] + explained_variance[1]) * 100
            
            interpretation = f"The first two principal components explain {total_var:.1f}% of the total variance in the data. "
            interpretation += f"PC1 accounts for {pc1_var:.1f}% and PC2 accounts for {pc2_var:.1f}% of the variance. "
            
            if total_var > 70:
                interpretation += "This suggests good dimensionality reduction with most variance captured."
            elif total_var > 50:
                interpretation += "This indicates moderate dimensionality reduction effectiveness."
            else:
                interpretation += "This suggests that additional components may be needed to capture more variance."
            
            return interpretation
        
        return "PCA analysis completed successfully."
    
    def _interpret_clustering_results(self, results: Dict[str, Any]) -> str:
        """Interpret clustering results"""
        if not results or "cluster_assignments" not in results:
            return "Clustering analysis completed but results interpretation unavailable."
        
        cluster_assignments = results["cluster_assignments"]
        if isinstance(cluster_assignments, dict):
            cluster_counts = {}
            for sample, cluster in cluster_assignments.items():
                cluster_counts[cluster] = cluster_counts.get(cluster, 0) + 1
            
            n_clusters = len(cluster_counts)
            total_samples = sum(cluster_counts.values())
            
            interpretation = f"The clustering analysis identified {n_clusters} distinct clusters from {total_samples} samples. "
            
            # Analyze cluster sizes
            cluster_sizes = list(cluster_counts.values())
            avg_size = sum(cluster_sizes) / len(cluster_sizes)
            
            if max(cluster_sizes) / min(cluster_sizes) > 3:
                interpretation += "The clusters show significant size variation, suggesting potential hierarchical structure."
            else:
                interpretation += "The clusters are relatively balanced in size."
            
            return interpretation
        
        return "Clustering analysis completed successfully."
    
    def _calculate_summary_statistics(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate summary statistics for the report"""
        stats = {
            "total_datasets": len(data["datasets"]),
            "total_analyses": len(data["analysis_results"]),
            "total_literature": len(data["literature_summaries"]),
            "total_genes": 0,
            "total_samples": 0,
            "analysis_types": {}
        }
        
        # Calculate dataset statistics
        for dataset in data["datasets"]:
            stats["total_genes"] += dataset.get("num_genes", 0)
            stats["total_samples"] += dataset.get("num_samples", 0)
        
        # Calculate analysis statistics
        for analysis in data["analysis_results"]:
            job_type = analysis.get("job_type", "unknown")
            stats["analysis_types"][job_type] = stats["analysis_types"].get(job_type, 0) + 1
        
        return stats
    
    async def _generate_report_content(self, data: Dict[str, Any], report_request: Dict[str, Any]) -> str:
        """Generate report content based on template"""
        template_name = report_request.get("template_name", "default")
        report_type = report_request["report_type"]
        
        # Select appropriate template
        template_file = f"{report_type}_{template_name}.html"
        
        # Try to load custom template, fallback to default
        try:
            template = self.jinja_env.get_template(template_file)
        except:
            # Create default template if not found
            template = self._create_default_template(report_type)
        
        # Render template with data
        content = template.render(
            data=data,
            options={
                "include_plots": report_request.get("include_plots", True),
                "include_statistics": report_request.get("include_statistics", True),
                "include_methodology": report_request.get("include_methodology", True)
            }
        )
        
        return content
    
    def _create_default_template(self, report_type: str) -> Template:
        """Create default template for report type"""
        if report_type == "analysis":
            template_content = """
            <!DOCTYPE html>
            <html>
            <head>
                <title>{{ data.report_metadata.title }}</title>
                <style>
                    body { font-family: Arial, sans-serif; margin: 20px; }
                    .header { text-align: center; margin-bottom: 30px; border-bottom: 2px solid #333; padding-bottom: 20px; }
                    .section { margin: 20px 0; }
                    .dataset { background: #f9f9f9; padding: 15px; margin: 10px 0; border-radius: 5px; }
                    .analysis { background: #e8f4f8; padding: 15px; margin: 10px 0; border-radius: 5px; }
                    .statistics { background: #fff3cd; padding: 15px; margin: 10px 0; border-radius: 5px; }
                    table { width: 100%; border-collapse: collapse; margin: 10px 0; }
                    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                    th { background-color: #f2f2f2; }
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>{{ data.report_metadata.title }}</h1>
                    <p>{{ data.report_metadata.description }}</p>
                    <p>Generated on: {{ data.report_metadata.generated_at }}</p>
                    <p>User: {{ data.user_info.name }} ({{ data.user_info.organization }})</p>
                </div>
                
                {% if options.include_statistics %}
                <div class="section">
                    <h2>Summary Statistics</h2>
                    <div class="statistics">
                        <p><strong>Total Datasets:</strong> {{ data.statistics.total_datasets }}</p>
                        <p><strong>Total Analyses:</strong> {{ data.statistics.total_analyses }}</p>
                        <p><strong>Total Genes:</strong> {{ data.statistics.total_genes }}</p>
                        <p><strong>Total Samples:</strong> {{ data.statistics.total_samples }}</p>
                    </div>
                </div>
                {% endif %}
                
                <div class="section">
                    <h2>Datasets</h2>
                    {% for dataset in data.datasets %}
                    <div class="dataset">
                        <h3>{{ dataset.name }}</h3>
                        <p><strong>Description:</strong> {{ dataset.description or 'No description' }}</p>
                        <p><strong>Organism:</strong> {{ dataset.organism or 'Not specified' }}</p>
                        <p><strong>Genes:</strong> {{ dataset.num_genes }}, <strong>Samples:</strong> {{ dataset.num_samples }}</p>
                        <p><strong>Quality Score:</strong> {{ dataset.quality_metrics.quality_score }}%</p>
                    </div>
                    {% endfor %}
                </div>
                
                <div class="section">
                    <h2>Analysis Results</h2>
                    {% for analysis in data.analysis_results %}
                    <div class="analysis">
                        <h3>{{ analysis.job_name }}</h3>
                        <p><strong>Type:</strong> {{ analysis.job_type }}</p>
                        <p><strong>Status:</strong> {{ analysis.status }}</p>
                        <p><strong>Interpretation:</strong> {{ analysis.interpretation }}</p>
                    </div>
                    {% endfor %}
                </div>
                
                {% if options.include_methodology %}
                <div class="section">
                    <h2>Methodology</h2>
                    <p>This report was generated using the BioIntel.AI platform for bioinformatics analysis.</p>
                </div>
                {% endif %}
            </body>
            </html>
            """
        else:
            # Default template for other report types
            template_content = """
            <!DOCTYPE html>
            <html>
            <head>
                <title>{{ data.report_metadata.title }}</title>
                <style>
                    body { font-family: Arial, sans-serif; margin: 20px; }
                    .header { text-align: center; margin-bottom: 30px; }
                    .section { margin: 20px 0; }
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>{{ data.report_metadata.title }}</h1>
                    <p>{{ data.report_metadata.description }}</p>
                    <p>Generated on: {{ data.report_metadata.generated_at }}</p>
                </div>
                
                <div class="section">
                    <h2>Report Content</h2>
                    <p>This is a {{ data.report_metadata.report_type }} report.</p>
                </div>
            </body>
            </html>
            """
        
        return Template(template_content)
    
    async def _generate_report_file(self, report: Report) -> str:
        """Generate report file in requested format"""
        file_path = self.reports_dir / f"report_{report.id}.{report.format_type}"
        
        if report.format_type == "pdf":
            # Generate PDF using WeasyPrint
            HTML(string=report.content).write_pdf(str(file_path))
        elif report.format_type == "docx":
            # Generate DOCX using python-docx
            await self._generate_docx(report.content, str(file_path))
        else:
            # Save as HTML
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(report.content)
        
        return str(file_path)
    
    async def _generate_docx(self, html_content: str, file_path: str):
        """Generate DOCX file from HTML content"""
        # Simple HTML to DOCX conversion
        # This is a basic implementation - could be enhanced with more sophisticated parsing
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('BioIntel.AI Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add basic content (simplified)
        paragraph = doc.add_paragraph()
        paragraph.add_run('This report was generated by BioIntel.AI platform.')
        
        # Add timestamp
        doc.add_paragraph(f'Generated on: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Save document
        doc.save(file_path)
    
    async def generate_report_file(self, report_id: int) -> str:
        """Generate report file on demand"""
        report = self.db.query(Report).filter(Report.id == report_id).first()
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        file_path = await self._generate_report_file(report)
        
        # Update report with file path
        report.file_path = file_path
        report.file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
        self.db.commit()
        
        return file_path
    
    async def list_templates(self) -> List[Dict[str, Any]]:
        """List available report templates"""
        templates = []
        
        # Default templates
        default_templates = [
            {
                "name": "default",
                "display_name": "Default Template",
                "description": "Standard report template with all sections",
                "report_types": ["analysis", "literature", "combined"]
            },
            {
                "name": "minimal",
                "display_name": "Minimal Template",
                "description": "Simplified report with essential information only",
                "report_types": ["analysis", "literature"]
            },
            {
                "name": "detailed",
                "display_name": "Detailed Template",
                "description": "Comprehensive report with extended analysis",
                "report_types": ["analysis", "combined"]
            }
        ]
        
        templates.extend(default_templates)
        
        # Scan for custom templates
        if self.templates_dir.exists():
            for template_file in self.templates_dir.glob("*.html"):
                if template_file.stem not in [t["name"] for t in default_templates]:
                    templates.append({
                        "name": template_file.stem,
                        "display_name": template_file.stem.replace("_", " ").title(),
                        "description": "Custom template",
                        "report_types": ["analysis", "literature", "combined"]
                    })
        
        return templates

# Global reports service instance
reports_service = ReportsService()